from __future__ import annotations
"""
Export des résultats : CSV, DOCX, PDF.
"""
import io
import re
from typing import Tuple, List, Optional
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from database import get_db
import models
from routers.auth import get_current_teacher
from config import TRIMESTRES_DATES
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)

router = APIRouter(prefix="/api/export", tags=["export"])


def _get_class_data(
    classe_id: str, trimestre: int, teacher: models.Teacher, db: Session
) -> Tuple[models.Classe, List[dict]]:
    classe = db.query(models.Classe).filter(
        models.Classe.id == classe_id,
        models.Classe.teacher_id == teacher.id,
    ).first()
    if not classe:
        raise HTTPException(status_code=404, detail="Classe introuvable")

    students = db.query(models.Student).filter(models.Student.classe_id == classe_id).all()
    rows = []
    for s in students:
        output = (
            db.query(models.LLMOutput)
            .filter(
                models.LLMOutput.student_id == s.id,
                models.LLMOutput.trimestre == trimestre,
            )
            .first()
        )
        rows.append(
            {
                "Nom": s.last_name,
                "Prénom": s.first_name,
                "Appréciation générale": output.general_appreciation if output else "",
                "Synthèse": output.synthesis if output else "",
                "Récompense suggérée": output.reward_suggestion if output else "",
                "Modifié manuellement": output.manually_edited if output else False,
            }
        )
    return classe, rows


@router.get("/{classe_id}/csv")
def export_csv(
    classe_id: str,
    trimestre: int,
    teacher: models.Teacher = Depends(get_current_teacher),
    db: Session = Depends(get_db),
):
    classe, rows = _get_class_data(classe_id, trimestre, teacher, db)
    df = pd.DataFrame(rows)
    buf = io.StringIO()
    df.to_csv(buf, index=False, encoding="utf-8-sig")
    buf.seek(0)
    filename = f"{classe.name}_T{trimestre}_resultats.csv".replace(" ", "_")
    return StreamingResponse(
        iter([buf.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{classe_id}/docx")
def export_docx(
    classe_id: str,
    trimestre: int,
    teacher: models.Teacher = Depends(get_current_teacher),
    db: Session = Depends(get_db),
):
    classe, rows = _get_class_data(classe_id, trimestre, teacher, db)
    doc = Document()
    doc.add_heading(f"{classe.name} — Trimestre {trimestre}", 0)

    for row in rows:
        doc.add_heading(f"{row['Nom']} {row['Prénom']}", level=2)
        doc.add_paragraph(f"Appréciation générale :\n{row['Appréciation générale']}")
        doc.add_paragraph(f"Synthèse :\n{row['Synthèse']}")
        doc.add_paragraph(f"Récompense : {row['Récompense suggérée']}")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{classe.name}_T{trimestre}_resultats.docx".replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ─── Helpers PDF ──────────────────────────────────────────────────────────────

def _fmt(val: Optional[float]) -> str:
    """Formate une moyenne en '14.36' ou '—' si absente."""
    if val is None:
        return "—"
    return f"{val:.2f}"


def _delta_str(delta: Optional[float]) -> str:
    """Formate un delta avec signe (+2.15 / -1.30 / —)."""
    if delta is None:
        return "—"
    sign = "+" if delta >= 0 else ""
    return f"{sign}{delta:.2f}"


def _delta_color(delta: Optional[float]) -> colors.Color:
    """
    Gradient de couleur pour la cellule Δ :
    - négatif → rouge (plus intense si |Δ| grand)
    - positif → vert (plus intense si |Δ| grand)
    - None ou zéro → blanc
    Seuil de saturation maximale : |Δ| = 3 points.
    """
    if delta is None or delta == 0:
        return colors.white
    intensity = min(abs(delta) / 3.0, 1.0)
    if delta < 0:
        return colors.Color(1.0, 1.0 - 0.7 * intensity, 1.0 - 0.7 * intensity)
    else:
        return colors.Color(1.0 - 0.7 * intensity, 1.0, 1.0 - 0.7 * intensity)


def _render_synthesis(synthesis: Optional[str], body: ParagraphStyle, bullet: ParagraphStyle) -> List:
    """
    Parse la synthèse structurée (Points forts / Axes d'amélioration / Alertes)
    et la retourne comme une liste d'éléments Paragraph avec labels en gras et bullet points.
    Si la synthèse n'est pas structurée, affiche une bullet par ligne.
    """
    if not synthesis:
        return [Paragraph("—", body)]

    pattern = re.compile(r"(Points forts|Axes?\s+d.amélioration|Alertes?)\s*:", re.IGNORECASE)
    parts = pattern.split(synthesis.strip())

    if len(parts) <= 1:
        # Pas de structure reconnue → une bullet par ligne
        result = []
        for line in synthesis.split("\n"):
            clean = line.strip().lstrip("-•").strip()
            if clean:
                result.append(Paragraph(f"• {clean}", body))
        return result or [Paragraph(synthesis.strip(), body)]

    result = []
    for i in range(1, len(parts) - 1, 2):
        label = parts[i].strip()
        content = parts[i + 1].strip() if i + 1 < len(parts) else ""
        result.append(Paragraph(f"<b>{label} :</b>", body))
        for line in re.split(r"\n|(?<=[.!?])\s+", content):
            clean = line.strip().lstrip("-•").strip()
            if clean:
                result.append(Paragraph(f"• {clean}", bullet))
    return result


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    """Crée un Paragraph en échappant les caractères XML si nécessaire."""
    safe = (text or "—").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


# ─── Export PDF ───────────────────────────────────────────────────────────────

@router.get("/{classe_id}/pdf")
def export_pdf(
    classe_id: str,
    trimestre: int,
    teacher: models.Teacher = Depends(get_current_teacher),
    db: Session = Depends(get_db),
):
    """
    Génère un PDF par élève avec :
    - Tableau de notes T_n vs T_{n-1} (gradient couleur sur le delta)
    - Bilan trimestre précédent (PP, mention, CE)
    - Incidents / Vie scolaire du trimestre
    - Appréciations LLM (général, synthèse, récompense)
    """
    classe = db.query(models.Classe).filter(
        models.Classe.id == classe_id,
        models.Classe.teacher_id == teacher.id,
    ).first()
    if not classe:
        raise HTTPException(status_code=404, detail="Classe introuvable")

    students = db.query(models.Student).filter(models.Student.classe_id == classe_id).all()

    # Dates du trimestre pour filtrer vie scolaire / sanctions
    dates_tn = TRIMESTRES_DATES.get(trimestre, {})
    debut_tn = dates_tn.get("debut", "")
    fin_tn = dates_tn.get("fin", "9999-12-31")
    prev_t = trimestre - 1

    # ── Styles ──────────────────────────────────────────────────────────────
    base = getSampleStyleSheet()

    h1 = ParagraphStyle("h1", fontSize=16, fontName="Helvetica-Bold", spaceAfter=6)
    h2 = ParagraphStyle("h2", fontSize=13, fontName="Helvetica-Bold", spaceBefore=10, spaceAfter=4)
    h3 = ParagraphStyle("h3", fontSize=11, fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=3, leftIndent=6)
    h4 = ParagraphStyle("h4", fontSize=10, fontName="Helvetica-BoldOblique", spaceBefore=6, spaceAfter=2, leftIndent=12)
    body = ParagraphStyle("body", fontSize=9, fontName="Helvetica", leading=12, leftIndent=12)
    label = ParagraphStyle("label", fontSize=9, fontName="Helvetica-Bold", leading=12, leftIndent=12)
    bullet = ParagraphStyle("bullet", fontSize=9, fontName="Helvetica", leading=12, leftIndent=24)
    cell = ParagraphStyle("cell", fontSize=8, fontName="Helvetica", leading=10)

    # ── Document ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.5 * cm, rightMargin=1.5 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )

    # Largeurs colonnes (total = 18 cm = usable width avec marges 1.5 cm)
    # Matière | Élève T_n | Classe T_n | Élève T_prev | Δ | Appréciation T_n | Appréciation T_prev
    col_widths = [3.0 * cm, 1.5 * cm, 1.5 * cm, 1.5 * cm, 1.5 * cm, 4.5 * cm, 4.5 * cm]

    story = []

    for student_idx, student in enumerate(students):
        if student_idx > 0:
            story.append(PageBreak())

        # ── Données élève ────────────────────────────────────────────────
        lines_tn = (
            db.query(models.BulletinLine)
            .filter(
                models.BulletinLine.student_id == student.id,
                models.BulletinLine.trimestre == trimestre,
            )
            .all()
        )
        lines_prev = (
            db.query(models.BulletinLine)
            .filter(
                models.BulletinLine.student_id == student.id,
                models.BulletinLine.trimestre == prev_t,
            )
            .all()
        ) if trimestre > 1 else []

        bilan_tn = next((l for l in lines_tn if l.subject == "BILAN"), None)
        bilan_prev = next((l for l in lines_prev if l.subject == "BILAN"), None)
        subjects_tn = [l for l in lines_tn if l.subject != "BILAN"]
        prev_by_subject = {l.subject: l for l in lines_prev if l.subject != "BILAN"}

        llm = (
            db.query(models.LLMOutput)
            .filter(
                models.LLMOutput.student_id == student.id,
                models.LLMOutput.trimestre == trimestre,
            )
            .first()
        )

        vie_scolaire = (
            db.query(models.VieScolaireEvent)
            .filter(
                models.VieScolaireEvent.student_id == student.id,
                models.VieScolaireEvent.date >= debut_tn,
                models.VieScolaireEvent.date <= fin_tn,
            )
            .order_by(models.VieScolaireEvent.date)
            .all()
        )
        sanctions = (
            db.query(models.SanctionEncouragement)
            .filter(
                models.SanctionEncouragement.student_id == student.id,
                models.SanctionEncouragement.date >= debut_tn,
                models.SanctionEncouragement.date <= fin_tn,
            )
            .order_by(models.SanctionEncouragement.date)
            .all()
        )

        # ── H1 : Prénom Nom ──────────────────────────────────────────────
        story.append(Paragraph(f"{student.first_name} {student.last_name}", h1))
        story.append(Spacer(1, 0.2 * cm))

        # ── H2 : Résultats ───────────────────────────────────────────────
        story.append(Paragraph("Résultats", h2))

        # En-tête du tableau
        t_prev_label = f"T{prev_t}" if trimestre > 1 else "T0"
        header_row = [
            _p("Matière", cell),
            _p(f"Élève T{trimestre}", cell),
            _p(f"Classe T{trimestre}", cell),
            _p(f"Élève {t_prev_label}", cell),
            _p(f"Δ T{trimestre}−{t_prev_label}", cell),
            _p(f"Appréciation T{trimestre}", cell),
            _p(f"Appréciation {t_prev_label}", cell),
        ]

        table_data = [header_row]
        # Styles de base : seront complétés cellule par cellule pour les deltas
        table_styles: List = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#D3D3D3")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("ALIGN", (1, 0), (4, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("LEFTPADDING", (0, 0), (-1, -1), 3),
            ("RIGHTPADDING", (0, 0), (-1, -1), 3),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]

        for row_idx, line in enumerate(subjects_tn, start=1):
            prev_line = prev_by_subject.get(line.subject)
            avg_tn = line.average
            avg_prev = prev_line.average if prev_line else None
            delta = (avg_tn - avg_prev) if (avg_tn is not None and avg_prev is not None) else None

            appr_prev_text = (prev_line.appreciation or "—") if prev_line else "—"
            row = [
                _p(line.subject, cell),
                _fmt(avg_tn),
                _fmt(line.average_class),
                _fmt(avg_prev),
                _delta_str(delta),
                _p(line.appreciation or "—", cell),
                _p(appr_prev_text, cell),
            ]
            table_data.append(row)
            table_styles.append(
                ("BACKGROUND", (4, row_idx), (4, row_idx), _delta_color(delta))
            )

        # Dernière ligne : Moyenne générale (depuis BILAN)
        bilan_row_idx = len(table_data)
        avg_tn_g = bilan_tn.average if bilan_tn else None
        avg_cls_tn = bilan_tn.average_class if bilan_tn else None
        avg_prev_g = bilan_prev.average if bilan_prev else None
        delta_g = (
            (avg_tn_g - avg_prev_g)
            if (avg_tn_g is not None and avg_prev_g is not None)
            else None
        )
        table_data.append([
            Paragraph("<b>Moyenne générale</b>", cell),
            _fmt(avg_tn_g),
            _fmt(avg_cls_tn),
            _fmt(avg_prev_g),
            _delta_str(delta_g),
            "",
            "",
        ])
        table_styles.extend([
            ("BACKGROUND", (0, bilan_row_idx), (3, bilan_row_idx), colors.HexColor("#EEEEEE")),
            ("BACKGROUND", (5, bilan_row_idx), (6, bilan_row_idx), colors.HexColor("#EEEEEE")),
            ("BACKGROUND", (4, bilan_row_idx), (4, bilan_row_idx), _delta_color(delta_g)),
            ("FONTNAME", (0, bilan_row_idx), (-1, bilan_row_idx), "Helvetica-Bold"),
        ])

        tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
        tbl.setStyle(TableStyle(table_styles))
        story.append(tbl)
        story.append(Spacer(1, 0.3 * cm))

        # ── H2 : Appréciations ───────────────────────────────────────────
        story.append(Paragraph("Appréciations", h2))

        # ── H3 : Trimestre précédent ─────────────────────────────────────
        if trimestre > 1 and bilan_prev:
            story.append(Paragraph(f"Trimestre précédent (T{prev_t})", h3))
            if bilan_prev.appreciation:
                story.append(Paragraph("Appréciation du professeur principal", label))
                story.append(Paragraph(bilan_prev.appreciation, body))
            story.append(Paragraph("Mention du conseil", label))
            story.append(Paragraph(bilan_prev.mention or "Pas de récompense / mention", body))
            if bilan_prev.appreciation_ce:
                story.append(Paragraph("Appréciation du chef d'établissement", label))
                story.append(Paragraph(bilan_prev.appreciation_ce, body))
            story.append(Spacer(1, 0.2 * cm))

        # ── H3 : Ce trimestre ────────────────────────────────────────────
        story.append(Paragraph(f"Ce trimestre (T{trimestre})", h3))

        # ── H4 : Incidents / Vie scolaire ────────────────────────────────
        abs_non_just = [e for e in vie_scolaire if e.event_type == "absence" and e.justifie is False]
        retards = [e for e in vie_scolaire if e.event_type == "retard"]

        if abs_non_just or retards or sanctions:
            story.append(Paragraph("Incidents / Vie scolaire", h4))

            if abs_non_just:
                story.append(Paragraph(f"Absences non justifiées ({len(abs_non_just)})", label))
                for e in abs_non_just:
                    parts = [e.display_date or e.date or ""]
                    if e.libelle:
                        parts.append(e.libelle)
                    if e.motif:
                        parts.append(f"— {e.motif}")
                    story.append(Paragraph("• " + " · ".join(p for p in parts if p), body))

            if retards:
                story.append(Paragraph(f"Retards ({len(retards)})", label))
                for e in retards:
                    parts = [e.display_date or e.date or ""]
                    if e.libelle:
                        parts.append(e.libelle)
                    if e.motif:
                        parts.append(f"— {e.motif}")
                    story.append(Paragraph("• " + " · ".join(p for p in parts if p), body))

            if sanctions:
                story.append(Paragraph(f"Sanctions / Encouragements ({len(sanctions)})", label))
                for s in sanctions:
                    parts = [s.display_date or s.date or ""]
                    if s.type_element:
                        parts.append(s.type_element)
                    if s.libelle:
                        parts.append(f"— {s.libelle}")
                    if s.motif:
                        parts.append(f"({s.motif})")
                    story.append(Paragraph("• " + " · ".join(p for p in parts if p), body))

            story.append(Spacer(1, 0.2 * cm))

        # ── H4 : Appréciation générale (LLM) ────────────────────────────
        story.append(Paragraph("Appréciation générale", h4))
        if llm:
            if llm.general_appreciation:
                story.append(Paragraph("Appréciation générale", label))
                story.append(Paragraph(llm.general_appreciation, body))
            if llm.synthesis:
                story.append(Paragraph("Synthèse", label))
                story.extend(_render_synthesis(llm.synthesis, body, bullet))
            if llm.reward_suggestion:
                story.append(Paragraph("Récompense suggérée", label))
                story.append(Paragraph(llm.reward_suggestion, body))
        else:
            story.append(Paragraph("Aucune appréciation générée.", body))

    doc.build(story)
    buf.seek(0)
    filename = f"{classe.name}_T{trimestre}_resultats.pdf".replace(" ", "_")
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
